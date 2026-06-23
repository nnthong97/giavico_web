from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "docs/beverage-formulation-service-build-spec.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, 16, True, "2E74B5")
    elif level == 2:
        set_run_font(run, 13, True, "2E74B5")
    else:
        set_run_font(run, 12, True, "1F4D78")
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, 11, False, "000000")
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, 11, False, "000000")
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_run_font(run, 11, False, "000000")
    return paragraph


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Definition"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "F2F4F7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, 10, True, "000000")
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, 10, False, "000000")
    set_table_width(table, [1.85, 4.65])
    return table


def add_checklist_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Check", "Expected Result", "Owner"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
    for check, expected, owner in rows:
        cells = table.add_row().cells
        cells[0].text = check
        cells[1].text = expected
        cells[2].text = owner
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, 10, row is table.rows[0], "000000")
    set_table_width(table, [1.75, 3.85, 0.9])
    return table


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
for style_name, size, color in [
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(12 if style_name != "Heading 3" else 8)
    style.paragraph_format.space_after = Pt(6 if style_name != "Heading 1" else 8)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)
run = title.add_run("Beverage Formulation Service Build Specification")
set_run_font(run, 22, True, "0B2545")

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
run = subtitle.add_run("Service handoff document for streaming Ollama-powered beverage R&D formula generation")
set_run_font(run, 11, False, "555555")

add_heading(doc, "Purpose", 1)
add_body(
    doc,
    "This document defines the service behavior, data contracts, integration points, and verification checks needed to build or maintain the beverage formulation generation service. It is based on the current Angular/Nx implementation in libs/data-access.",
)

add_heading(doc, "Service Scope", 1)
for item in [
    "Generate a structured beverage formula from target parameters and operational constraints.",
    "Support both complete-response generation and real-time streaming generation.",
    "Normalize and validate formula output before the UI treats it as final.",
    "Expose saved formula metadata for list and detail views.",
    "Preserve local-first operation against an Ollama endpoint at http://localhost:11434/api/generate.",
]:
    add_bullet(doc, item)

add_heading(doc, "Primary Service", 1)
add_key_value_table(
    doc,
    [
        ("Class", "OllamaFormulationService"),
        ("Location", "libs/data-access/src/lib/services/ollama-formulation.service.ts"),
        ("Default endpoint", "http://localhost:11434/api/generate"),
        ("Default model", "gemma4"),
        ("Synchronous method", "generateFormula(input, historicalData?, model?, apiUrl?)"),
        ("Streaming method", "generateFormulaStream(input, historicalData?, model?, apiUrl?)"),
    ],
)

add_heading(doc, "Input Contract", 1)
add_key_value_table(
    doc,
    [
        ("drinkName", "Human-readable product name used in the prompt and saved formula title."),
        ("marketDestination", "Target market or region for regulatory and sensory assumptions."),
        ("targetBrix", "Target sugar/solids level in degrees Brix."),
        ("isAcidified", "Whether the formula must satisfy acidified beverage constraints below pH 4.6."),
        ("regionalRestrictions", "Array of regulatory, ingredient, or label restrictions."),
        ("productionArea", "Factory line, equipment, or production context."),
        ("customerSpecification", "Customer-facing product requirements and sensory constraints."),
        ("baselineBOM", "Baseline bill of materials identifier or pattern reference."),
    ],
)

add_heading(doc, "Output Contract", 1)
add_key_value_table(
    doc,
    [
        ("ingredients", "Array of raw material rows with rawMaterialKey, massPercentage, and costProjection."),
        ("varianceAnalysis", "Markdown-like explanation of differences, swaps, and sensory impact."),
        ("stabilityAlerts", "Array of safety, processing, regulatory, or mass-balance warnings."),
        ("savedFormula", "Session-level saved record with id, name, summary, savedAt, and formula."),
    ],
)

add_heading(doc, "Streaming Behavior", 1)
add_body(doc, "The streaming method should emit two event types through an Observable.")
add_key_value_table(
    doc,
    [
        ("progress", "Emitted whenever a streamed Ollama response fragment arrives. The payload is the accumulated partial JSON string."),
        ("complete", "Emitted once the stream ends, the accumulated JSON parses successfully, and formula validation/normalization completes."),
    ],
)
for step in [
    "Create an AbortController per subscription so cancellation aborts the underlying fetch.",
    "POST an Ollama request with format: json and stream: true.",
    "Read response.body using a ReadableStream reader and decode chunks with TextDecoder.",
    "Split incoming text on newline boundaries because Ollama streams newline-delimited JSON objects.",
    "Parse each line as an OllamaResponse and append the response field to the accumulated output.",
    "Emit progress with the accumulated partial response after every non-empty response fragment.",
    "At stream completion, parse the accumulated JSON as BeverageFormula, normalize it, emit complete, and close the observer.",
]:
    add_number(doc, step)

add_heading(doc, "Validation Rules", 1)
for item in [
    "ingredients must exist and be an array.",
    "varianceAnalysis defaults to 'No variance analysis provided by the LLM model.' when missing or invalid.",
    "stabilityAlerts defaults to an empty array when missing or invalid.",
    "Ingredient mass percentages are summed; when the total differs from 100 by more than 0.5, add a mass-balance warning to stabilityAlerts.",
    "Parsing failures should surface as explicit errors that distinguish full response parsing from streamed response parsing.",
]:
    add_bullet(doc, item)

add_heading(doc, "NgRx Integration", 1)
add_key_value_table(
    doc,
    [
        ("requestAIGeneration", "Starts generation, stores input, clears old formula and streamingResponse, sets loading true."),
        ("generationProgress", "Stores the latest accumulated streamed JSON in streamingResponse."),
        ("generationSuccess", "Stores final formula, clears streamingResponse, sets loading false."),
        ("generationFailure", "Stores error and sets loading false."),
        ("saveFormula", "Prepends a saved formula record to savedFormulas."),
    ],
)

add_heading(doc, "UI Responsibilities", 1)
for item in [
    "Target Formulation Matrix must include Drink Name, market, Brix, acidification, restrictions, production area, customer specification, and baseline BOM.",
    "During generation, show the existing loading state plus the Live Formula Stream block when streamingResponse is non-empty.",
    "After completion, hide the stream and show the existing ingredients matrix, variance analysis, stability alerts, and Save button.",
    "Saved Formulas list items should show drink name, formula summary, ingredient count, total cost per liter, total mass percentage, and saved datetime.",
    "Clicking a saved formula item should navigate to /formulas/:id and render the detail page from store state.",
]:
    add_bullet(doc, item)

add_heading(doc, "Error Handling", 1)
for item in [
    "Network errors should produce a clear connection message for the local Ollama endpoint.",
    "Timeouts should mention the generation timeout limit and suggest local Ollama load as a likely cause.",
    "HTTP non-OK responses should include the returned status code.",
    "Missing response.body should be treated as a streaming capability failure.",
    "Invalid streamed JSON should include a parse failure message and stop loading.",
]:
    add_bullet(doc, item)

add_heading(doc, "Implementation Checklist", 1)
add_checklist_table(
    doc,
    [
        ("Create service method generateFormulaStream", "Observable emits progress and complete events.", "Data access"),
        ("Wire effects to streaming method", "Progress and success actions update state correctly.", "NgRx"),
        ("Expose streamingResponse selector", "UI can subscribe via async pipe.", "NgRx"),
        ("Render live stream in output panel", "Partial JSON appears before final formula.", "UI"),
        ("Add Drink Name to form and prompt", "Saved title and model context use product name.", "UI/service"),
        ("Add saved formula detail route", "Clicking list item opens /formulas/:id.", "Routing"),
    ],
)

add_heading(doc, "Verification Checklist", 1)
add_checklist_table(
    doc,
    [
        ("TypeScript compile", "npx tsc -p apps/dashboard/tsconfig.app.json --noEmit passes.", "Engineer"),
        ("Angular template compile", "npx ngc -p apps/dashboard/tsconfig.app.json --noEmit passes.", "Engineer"),
        ("Lint", "npx nx lint dashboard --skip-nx-cache passes.", "Engineer"),
        ("Streaming runtime", "Live Formula Stream appears while Ollama is generating.", "Engineer"),
        ("Final formula", "Ingredients matrix and Save button render after completion.", "Engineer"),
        ("Saved item", "Saved list displays title, summary, ingredients, cost, mass percentage, datetime.", "Engineer"),
        ("Detail navigation", "Clicking saved item navigates to /formulas/:id with detail data.", "Engineer"),
    ],
)

add_heading(doc, "Known Build Caveat", 1)
add_body(
    doc,
    "In the current environment, nx build dashboard has repeatedly failed with an esbuild deadlock. TypeScript, Angular template compilation, linting, and browser runtime checks pass. Treat the build issue as an environment/toolchain blocker unless reproduced with a normal compiler diagnostic.",
)

doc.save(OUTPUT)
print(OUTPUT)
